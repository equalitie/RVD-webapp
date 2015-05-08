# -*- coding: utf-8 -*-

import re
import time
import datetime
import docx
import xlrd

import utils
from rvd.models import *

# Constants

# Identifiers for types of legacy documents

ORG1_DOCX = 1
ORG2_DOCX = 2
EXCEL_DOC = 4


# Private - Should not be called externally

def _parse_org1_docx_report(doc):
    report = ''
    for p in doc.paragraphs:
        para = p.text.strip()
        if len(para) > 0:
            report += para + '\n'
    return report


def _parse_org1_docx_events(doc):
    events = []
    table = doc.tables[0]
    # First row contains table headers
    for i in range(1, len(table.rows)):
        row = table.row_cells(i)
        names = row[2].text.split(',')
        # Organisation names are provided in parentheses at the end of the list of names
        orgindex = names[-1].find('(')
        if orgindex >= 0:
            organisation = names[-1][orgindex + 1: names[-1].find(')')]
            names[-1] = names[-1][:orgindex]
        for name in names:
            parsed = {'actor': {}, 'location': {}, 'source': {}}
            parsed['actor']['name'] = name.strip()
            parsed['actor']['organisations'] = [organisation]
            parsed['report_date'] = time.strptime(row[0].text.strip(), '%d.%m.%y')
            parsed['report_date'] = utils.to_datetime(parsed['report_date'])
            parsed['location']['name'] = row[1].text.strip()
            # TODO
            # Think of how this field fits into the schema we have.
            # parsed['type'] = row[3].text.strip()
            parsed['source']['name'] = row[4].text.strip()
            events.append(parsed)
    return events

def _org1_events_to_model(events):
    '''Convert a collection of events into instances of the Event model'''
    for i in range(len(events)):
        event = events[i]
        actor = Actor(name=event['actor'])
        del event['actor']
        geocoded = utils.geocodes(event['location']['name'], include_importance=True)
        geocoded = utils.max_by(geocoded, lambda gc: gc.importance)
        location = Location(name=geocoded['name'],
            latitude=geocoded['latitude'], longitude=geocoded['latitude'])
        del event['location']
        events[i] = Event(**event)
        events[i].actor = actor
        events[i].location = location
    return events

def _org1_report_to_model(report, events):
    '''Convert a report text into an instance of the Report model'''
    report = Report(text=report)
    for i in range(events):
        events[i].report_id = report
    report.events = events
    return report

def _parse_org1_docx(_file):
    doc = docx.Document(_file)
    report = _parse_org1_docx_report(doc)
    events = _parse_org1_docx_events(doc)
    events = _org1_events_to_model(events)
    report = _org1_report_to_model(report, events)
    return {
        'report': report,
        'events': events
    }


# Parse the report part of a document before the list of event descriptions
def _parse_org2_docx_report(doc):
    report = ''
    for p in doc.paragraphs:
        if re.match('\d+\)\s+(n|N)ombre', p.text) is not None:
            break
        report += p.text + '\n'
    return report


# parse the list of event descriptions
def _parse_org2_docx_events(doc):
    events = []
    for i, p in enumerate(doc.paragraphs):
        if re.match('\d+\)\s+(n|N)ombre', p.text) is not None:
            parsed = {'actor': {}, 'location': {}, 'source': {}}
            parsed['actor']['name'] = ' '.join(doc.paragraphs[i].text.split(':')[1:]).strip()
            parsed['actor']['organisation'] = ' '.join(doc.paragraphs[i + 1].text.split(':')[1:]).strip()
            parsed['location']['name'] = ' '.join(doc.paragraphs[i + 2].text.split(':')[1:]).strip()
            # TODO
            # Decide if we want to include this extra field that does not follow
            # the schema description
            # parsed['notes'] = ' '.join(doc.paragraphs[i + 3].text.split(':')[1:]).strip()
            parsed['source']['name'] = ' '.join(doc.paragraphs[i + 4].text.split(':')[1:]).strip()
            parsed['detention_date'] = ' '.join(doc.paragraphs[i + 5].text.split(':')[1:]).strip()
            parsed['report_date'] = ' '.join(doc.paragraphs[i + 6].text.split(':')[1:]).strip()
            if parsed['detention_date']:
                parsed['detention_date'] = parsed['detention_date'].replace(' ', '').replace('.', '')
                try:
                    parsed['detention_date'] = time.strptime(parsed['detention_date'], '%d-%m-%Y')
                except:
                    parsed['detention_date'] = time.strptime(parsed['detention_date'], '%d-%m-%y')
                parsed['detention_date'] = utils.to_datetime(parsed['detention_date'])
            if parsed['report_date']:
                parsed['report_date'] = parsed['report_date'].replace(' ', '').replace('.', '')
                try:
                    parsed['report_date'] = time.strptime(parsed['report_date'], '%d-%m-%Y')
                except:
                    parsed['report_date'] = time.strptime(parsed['report_date'], '%d-%m-%y')
                parsed['report_date'] = utils.to_datetime(parsed['report_date'])
            events.append(parsed)
    return events

def _org2_events_to_model(events):
    '''Convert parsed event data to model instances'''
    for i in range(len(events)):
        event = events[i]
        actor = event['actor']
        organisation = actor['organisation']
        actor = Actor(name=actor['name'])
        organisation = Organisation(name=organisation)
        actor.organisations = [organisation]
        location = event['location']
        geocodes = utils.geocodes(location, include_importance=True)
        location = utils.max_by(geocodes, lambda gc: gc['importance'])
        location = Location(name=location['name'],
            latitude=location['latitude'], longitude=location['longitude'])
        source = event['source']
        source = Source(name=source['name'])
        del event['actor']
        del event['location']
        del event['source']
        event['victims'] = [actor]
        event['locations'] = [location]
        event['sources']  = [source]
        events[i] = Event(**event)
    return events

def _org2_report_to_model(report, events):
    '''Convert parsed report data to a model instnace'''
    report = Report(text=report)
    for i in range(len(events)):
        events[i].report_id = report.id
    report.events = events
    return report

def _parse_org2_docx(_file):
    doc = docx.Document(_file)
    report = _parse_org2_docx_report(doc)
    events = _parse_org2_docx_events(doc)
    events = _org2_events_to_model(events)
    report = _org2_report_to_model(report, events)
    return {
        'report': _parse_org2_docx_report(doc),
        'events': _parse_org2_docx_events(doc)
    }


# This function will do generic parsing of the content in an excel workbook
# but we will leave it to other handler functions to do entity-specific manipulation
def _parse_excel_template(filename):
    book = xlrd.open_workbook(filename)
    # We assume everything is on the same sheet
    sheet = book.sheet_by_index(0)
    entities = {}
    current_entity_fields = []
    current_entity_name = ''
    # The first row contains just the table name, so we skip it
    for row_num in range(1, sheet.nrows):
        # When we encounter a row like ['Event', 'Title', 'Description', ...]
        # We are looking at the outline of a new entity type (in example, it's 'Event')
        if sheet.row(row_num)[0].value != '':
            # Before transitioning to the next entity, apply any special handlers
            # We might want to use to do extra work with the current entity's data
            if current_entity_name != '':
                handler = _post_collection_handlers.get(current_entity_name, _id)
                entities[current_entity_name] = handler(entities[current_entity_name], book)
            entity_name = sheet.row(row_num)[0].value
            entities[entity_name] = []
            current_entity_name = entity_name
            current_entity_fields = [col.value for col in sheet.row_slice(row_num, 1)]
            continue
        values = [col.value for col in sheet.row_slice(row_num, 1)]
        new_instance = {}
        for i in range(len(current_entity_fields)):
            field_name = current_entity_fields[i]
            new_instance[field_name] = values[i]
        entities[current_entity_name].append(new_instance)
    handler = _post_collection_handlers.get(current_entity_name, _id)
    entities[current_entity_name] = handler(entities[current_entity_name], book)
    return entities


def _parse_error(stream):
    return {'error': 'No parser exists for the provided filetype.'}


def _id(x, *args, **kwargs):
    return x


def _excel_parse_date(datefloat, workbook):
    return datetime.datetime(*xlrd.xldate_as_tuple(datefloat, workbook.datemode))


def _event_entity_handler(events, wb):
    for i in range(len(events)):
        events[i]['Date of detention'] = _excel_parse_date(events[i]['Date of detention'], wb)
        events[i]['Date of release'] = _excel_parse_date(events[i]['Date of release'], wb)
        events[i]['Date of report'] = _excel_parse_date(events[i]['Date of report'], wb)
        locations = utils.geocodes(events[i]['Location'], include_importance=True)
        best_location = utils.max_by(locations, lambda loc: loc['importance'])
        events[i]['Location'] = best_location
    return events


def _actor_entity_handler(actors, wb):
    for i in range(len(actors)):
        actors[i]['Date of birth'] = _excel_parse_date(actors[i]['Date of birth'], wb)
        orgs = actors[i]['Organisations']
        actors[i]['Organisations'] = [s.strip() for s in orgs.split(',')]
        try:
            actors[i]['Age'] = int(actors[i]['Age'])
        except:
            pass
        locations = utils.geocodes(actors[i]['Address'], include_importance=True)
        best_location = utils.max_by(locations, lambda loc: loc['importance'])
        actors[i]['Address'] = best_location
    return actors


def _prison_entity_handler(prisons, wb):
    for i in range(len(prisons)):
        locations = utils.geocodes(prisons[i]['Location'], include_importance=True)
        best_location = utils.max_by(locations, lambda loc: loc['importance'])
        prisons[i]['Location'] = best_location
    return prisons


def _source_entity_handler(sources, wb):
    for i in range(len(sources)):
        orgs = sources[i]['Organisations']
        sources[i]['Organisations'] = [s.strip() for s in orgs.split(',')]
    return sources


def _organisation_entity_handler(organisations, wb):
    # Each organisation can have multiple locations
    for i in range(len(organisations)):
        geocoded = []
        for location in organisations[i]['Locations']:
            locations = utils.geocodes(location, include_importance=True)
            best_location = utils.max_by(locations, lambda loc: loc['importance'])
            geocoded.append(best_location)
        organisations[i]['Locations'] = geocoded
    return organisations


_post_collection_handlers = {
    'Events': _event_entity_handler,
    'Actors': _actor_entity_handler,
    'Prisons': _prison_entity_handler,
    'Sources': _source_entity_handler,
    'Organisations': _organisation_entity_handler
}

_parsers = {
    ORG1_DOCX: _parse_org1_docx,
    ORG2_DOCX: _parse_org2_docx,
    EXCEL_DOC: _parse_excel_template
}


# Public

def parse(stream, filetype):
    '''Parse the contents of a report file with one of the accepted filetypes into a dict'''
    parse = _parsers.get(filetype, _parse_error)
    return parse(stream)
