# -*- coding: utf-8 -*-

from flask_babel import lazy_gettext as ___

import re
import time
import datetime
import docx
import xlrd

import utils
from models import *
import entities

# Constants

# Identifiers for types of legacy documents

ORG1_DOCX = 1
ORG2_DOCX = 2
EXCEL_DOC = 4

# Special values

# Values compared to these should always be converted to lowercase first.
YES = ___('yes')
NO = ___('no')
MALE = ___('male')
FEMALE = ___('female')

# Entity names, as presented in the excel document template

ACTIONS = u'Actions'
ACTORS = u'Actors'
REPORTS = u'Reports'
EVENTS = u'Events'
INTERNATIONAL_AUTHORITIES = u'International Authorities'
LOCATIONS = u'Locations'
ORGANISATIONS = u'Organisations'
PRISONS = u'Prisons'
PRISON_TYPES = u'Prison Types'
PROFESSIONS = u'Professions'
RELEASE_TYPES = u'Release Types'
RIGHTS_VIOLATIONS = u'Rights Violations'
SOURCES = u'Sources'
STATE_AUTHORITIES = u'State Authorities'

# Private - Should not be called externally

def log_error(msg):
    print msg
    # Log to a file later?

def _excel_parse_date(datefloat, workbook):
    return datetime.datetime(*xlrd.xldate_as_tuple(datefloat, workbook.datemode))

def _parse_org1_docx_report(doc):
    report = ''
    for p in doc.paragraphs:
        para = p.text.strip()
        if len(para) > 0:
            report += para + '\n'
    return report

# Parse organisation 1's docx documents.
# Does not accept the excel_workbook argument (None, here)
def _parse_org1_docx_events(doc):
    events = []
    table = doc.tables[0]
    # First row contains table headers
    for i in range(1, len(table.rows)):
        row = table.row_cells(i)
        names = row[2].text.split(',')
        # Organisation names are provided in parentheses at the end of the list of names
        orgindex = names[-1].find('(')
        if orgindex >= 0:
            organisation = names[-1][orgindex + 1: names[-1].find(')')]
            names[-1] = names[-1][:orgindex]
        for name in names:
            parsed = {'actor': {}, 'location': {}, 'source': {}}
            parsed['actor']['name'] = name.strip()
            parsed['actor']['organisations'] = [organisation]
            parsed['report_date'] = time.strptime(row[0].text.strip(), '%d.%m.%y')
            parsed['report_date'] = utils.to_datetime(parsed['report_date'])
            parsed['location']['name'] = row[1].text.strip()
            # TODO
            # Think of how this field fits into the schema we have.
            # parsed['type'] = row[3].text.strip()
            parsed['source']['name'] = row[4].text.strip()
            events.append(parsed)
    return events

def _org1_events_to_model(events):
    '''Convert a collection of events into instances of the Event model'''
    for i in range(len(events)):
        event = events[i]
        actor = Actor(name=event['actor']['name'])
        del events[i]['actor']
        source = Source(name=event['source']['name'])
        del events[i]['source']
        location = _get_location(event['location']['name'])
        del events[i]['location']
        events[i]['locations'] = [location]
        events[i]['victims'] = [actor]
        events[i]['sources'] = [source]
        events[i] = Event(**events[i])
    return events

def _org1_report_to_model(report, events):
    '''Convert a report text into an instance of the Report model'''
    report = Report(text=report)
    for i in range(len(events)):
        events[i].report_id = report
    report.events = events
    return report

def _parse_org1_docx(_file):
    doc = docx.Document(_file)
    report = _parse_org1_docx_report(doc)
    events = _parse_org1_docx_events(doc)
    events = _org1_events_to_model(events)
    report = _org1_report_to_model(report, events)
    session.add(report)
    for i, event in enumerate(events):
        event.title = 'Evento ' + str(i)
        session.add_all(event.locations)
        session.add_all(event.victims)
        session.add_all(event.sources)
        session.add(event)
    session.commit()
    return {
        'report': report,
        'events': events
    }


# Parse the report part of a document before the list of event descriptions
def _parse_org2_docx_report(doc):
    report = ''
    for p in doc.paragraphs:
        if re.match('\d+\)\s+(n|N)ombre', p.text) is not None:
            break
        report += p.text + '\n'
    return report


# Parse organisation 2's docx documents.
# Does not accept the excel_workbook argument (None, here)
def _parse_org2_docx_events(doc):
    def clean_text(pg):
        return ' '.join(pg.text.split(':')[1:]).strip()
    events = []
    for i, p in enumerate(doc.paragraphs):
        if re.match('\d+\)\s+(n|N)ombre', p.text) is not None:
            parsed = {'actor': {}, 'location': {}, 'prison': {}, 'source': {}}
            parsed['actor']['name'] = clean_text(doc.paragraphs[i])
            parsed['actor']['organisation'] = clean_text(doc.paragraphs[i + 1])
            parsed['location']['name'] = clean_text(doc.paragraphs[i + 2])
            parsed['consequences'] = clean_text(doc.paragraphs[i + 3])
            parsed['charges'] = clean_text(doc.paragraphs[i + 4])
            parsed['prison']['name'] = clean_text(doc.paragraphs[i + 5])
            parsed['actor']['telephone'] = clean_text(doc.paragraphs[i + 6])
            parsed['source']['name'] = clean_text(doc.paragraphs[i + 8])
            parsed['report_date'] = clean_text(doc.paragraphs[i + 9])
            # If a report date was supplied, parse it out and convert it to a datetime object.
            if len(parsed['report_date']) > 0:
                parsed['report_date'] = parsed['report_date'].replace('.', '')
                try:
                    parsed['report_date'] = time.strptime(parsed['report_date'], '%d-%m-%Y')
                except:
                    parsed['report_date'] = time.strptime(parsed['report_date'], '%d-%m-%y')
                parsed['report_date'] = utils.to_datetime(parsed['report_date'])
            else:
                del parsed['report_date']
            events.append(parsed)
    return events

def _org2_events_to_model(events):
    '''Convert parsed event data to model instances'''
    for i in range(len(events)):
        event = events[i]
        actor = event['actor']
        organisation = actor['organisation']
        actor = Actor(name=actor['name'])
        organisation = Organisation(name=organisation)
        actor.organisations = [organisation]
        location = event['location']['name']
        if location.endswith('.'):
            location = location[:-1]
        location = _get_location(location)
        event['locations'] = [location]
        prison = event['prison']
        prison = Prison(name=prison['name'])
        source = event['source']
        source = Source(name=source['name'])
        del event['actor']
        del event['location']
        del event['prison']
        del event['source']
        event['victims'] = [actor]
        event['prisons'] = [prison]
        event['sources']  = [source]
        events[i] = Event(**event)
    return events

def _org2_report_to_model(report, events):
    '''Convert parsed report data to a model instnace'''
    report = Report(text=report)
    for i in range(len(events)):
        events[i].report_id = report.id
    report.events = events
    return report

def _parse_org2_docx(_file):
    doc = docx.Document(_file)
    report = _parse_org2_docx_report(doc)
    events = _parse_org2_docx_events(doc)
    events = _org2_events_to_model(events)
    report = _org2_report_to_model(report, events)
    session.add(report)
    for i, event in enumerate(events):
        event.title = 'Evento ' + str(i)
        session.add_all(event.locations)
        session.add_all(event.victims)
        session.add_all(event.prisons)
        session.add_all(event.sources)
        session.add(event)
    session.commit()
    return {
        'report': report,
        'events': events
    }

def _get_location(loc_name):
    location = session.query(Location).filter_by(name=loc_name).first()
    if location is not None:
        return location
    # Create a new Location based on the data pulled from OSM if the location
    # isn't already in the database
    locations = utils.geocodes(loc_name, include_importance=True)
    location = utils.max_by(locations, lambda l: l['importance'])
    location = Location(name=location['name'],
        longitude=location['longitude'], latitude=location['latitude'])
    return location

def _parse_excel_template(filename):
    book = xlrd.open_workbook(filename)
    sheet = book.sheet_by_index(0)
    cur_row = 0
    cur_event = None
    while cur_row < sheet.nrows:
        row = sheet.row(cur_row)
        #print u'ROW [' + u', '.join([str(row[i].value) for i in range(len(row))]) + u']'
        if u'Titulo' in row[1].value:
            cur_row += 1
            row = sheet.row(cur_row)
            #print u'EVENT [' + u', '.join([str(row[i].value) for i in range(len(row))]) + u']'
            cur_event = Event(**{
                'title': row[1].value,
                'description': row[2].value,
                'charges': row[3].value,
                'consequences': row[4].value,
                'event_start': _excel_parse_date(float(row[5].value), book),
                'event_end': _excel_parse_date(float(row[6].value), book),
                'report_date': _excel_parse_date(float(row[7].value), book),
                'victim_is_complainant': row[8].value.lower() == YES,
                'allow_storage': row[9].value.lower() == YES,
                'allow_publishing': row[10].value.lower() == YES,
                'allow_representation': row[11].value.lower() == YES,
                'data_is_sensitive': row[12].value.lower() == YES,
                'release_types': [session.query(ReleaseType).filter_by(type_code=int(row[13].value)).first()],
                'locations': [_get_location(row[14].value)],
                'prisons': [session.query(Prison).filter_by(name=row[15].value).first()],
                'event_types': [session.query(EventType).filter_by(name=row[16].value).first()],
                'owner': session.query(User).filter_by(is_admin=1).first()
            })
            cur_event.victims = []
            cur_event.witnesses = []
            cur_event.perpetrators = []
            cur_event.actions = []
            cur_event.sources = []
        elif row[1].value == u'Actors':
            #print u'Found Actors row; current event.title = ' + cur_event.title
            cur_row += 1
            row = sheet.row(cur_row)
            # Rely on the columns under an entity name and before the next being blank
            while row[1].value == '' and cur_row < sheet.nrows:
                #print u'ACTOR [' + u', '.join([str(row[i].value) for i in range(len(row))]) + u']'
                _type = row[2].value
                actor = Actor(**{
                    'name': row[3].value,
                    'organisations': [session.query(Organisation).filter_by(name=row[4].value).first()],
                    'birth_date': _excel_parse_date(float(row[5].value), book),
                    'telephone': row[6].value,
                    'address': row[7].value,
                    'gender': row[8].value,
                    'is_activist': row[9].value.lower() == YES,
                    'activist_info': row[10].value,
                    'professions': [session.query(Profession).filter_by(name=row[11].value).first()],
                    'owner': session.query(User).filter_by(is_admin=1).first()
                })
                #print u'Parsed actor<' + u', '.join(map(str, [
                    #actor.name, actor.birth_date, actor.telephone, actor.gender,
                    #actor.is_activist, actor.organisations[0].name, actor.professions[0].name])) + u'>'
                if _type == 'Witness':
                    cur_event.witnesses.append(actor)
                elif _type == 'Victim':
                    cur_event.victims.append(actor)
                elif _type == 'Perpetrator':
                    cur_event.perpetrators.append(actor)
                if (cur_row + 1) < sheet.nrows and sheet.row(cur_row + 1)[1].value == '':
                    cur_row += 1
                    row = sheet.row(cur_row)
                else:
                    break
        elif row[1].value == u'Sources':
            #print u'Found sources row'
            cur_row += 1
            row = sheet.row(cur_row)
            while row[1].value == '' and cur_row < sheet.nrows:
                #print u'SOURCE [' + u', '.join([str(row[i].value) for i in range(len(row))]) + u']'
                row = sheet.row(cur_row)
                source = Source(**{
                    'name': row[2].value,
                    'organisations': [session.query(Organisation).filter_by(name=row[3].value).first()]
                })
                #print u'Parsed source<' + u', '.join(map(str, [
                    #source.name, source.organisations[0].name])) + u'>'
                cur_event.sources.append(source)
                if (cur_row + 1) < sheet.nrows and sheet.row(cur_row + 1)[1].value == '':
                    cur_row += 1
                    row = sheet.row(cur_row)
                else:
                    break
        elif row[1].value == u'Report':
            #print 'Found report row'
            report = Report(**{
                'text': row[2].value,
                'events': [cur_event]
            })
            session.add(report)
            #print u'Parsed report: ' + report.text
        cur_row += 1


    if cur_event is not None:
        session.add_all(cur_event.witnesses)
        session.add_all(cur_event.victims)
        session.add_all(cur_event.perpetrators)
        session.add_all(cur_event.actions)
        session.add_all(cur_event.sources)
        session.add(cur_event)

    session.commit()

'''
# This function will do generic parsing of the content in an excel workbook
# but we will leave it to other handler functions to do entity-specific manipulation
def _parse_excel_template(filename):
    book = xlrd.open_workbook(filename)
    # We assume everything is on the same sheet
    sheet = book.sheet_by_index(0)
    entities = {}
    current_entity_fields = []
    current_entity_name = ''
    for row_num in range(sheet.nrows):
        # When we encounter a row like ['Event', 'Title', 'Description', ...]
        # We are looking at the outline of a new entity type (in example, it's 'Event')
        if sheet.row(row_num)[0].value != '':
            entity_name = sheet.row(row_num)[0].value
            entities[entity_name] = []
            current_entity_name = entity_name
            current_entity_fields = [col.value for col in sheet.row_slice(row_num, 1)]
            continue
        values = [col.value for col in sheet.row_slice(row_num, 1)]
        new_instance = {}
        for i in range(len(current_entity_fields)):
            field_name = current_entity_fields[i]
            new_instance[field_name] = values[i]
        if '' in new_instance:
            del new_instance['']
        entities[current_entity_name].append(new_instance)
    entities['finished'] = [] 
    # Call on each entity's respective model instance builder.
    entity_names = entities.keys()
    for entity_name in entity_names:
        #print entities['finished'], entity_name
        if entity_name not in entities['finished']:
            handler = _post_collection_handlers.get(entity_name, _id)
            entities = handler(entities, book)
            entities['finished'].append(entity_name)
    del entities['finished']
    return entities



def _id(x, *args, **kwargs):
    return x



'''
# Given a list of strings of integers that should be indices into an array,
# produce a list of converted integers. A string error message will be left
# in the place of a value that could not be parsed. Indeces are 1-indexed.
def _make_indices(indices, array_len):
    for i in range(len(indices)):
        try:
            v = int(indices[i])
        except ValueError:
            indices[i] = '{0} not a valid integer.'.format(indices[i])
            continue
        if v < 1 or v > array_len:
            indices[i] = 'Index ({0}) must be between 1 and {1} inclusive'.format(
                v, array_len)
        else:
            indices[i] = v - 1
    return indices

# Using a list of indices (or error strings) as generated by _make_indices,
# take items from a given array with the specified indices in order.
def _take(array, indices):
    pulled = []
    for index in indices:
        if isinstance(index, int):
            pulled.append(array[index])
        else: # index is an error string
            log_error(index)
    return pulled

# A convenience function to convert a string listing indices (e.g. "1, 2, 4")
# into the integer equivalent, 0-indexed, array ([0, 1, 3]) and then
# Extract elements from the provided array being indexed into and outputting
# any errors that occur.
def _from_indices(ind_list, other_entities):
    # If a single index value is provided, it will be converted to a float.
    # So it is necessary to convert it back to a string-int.
    if isinstance(ind_list, float):
        ind_list = str(int(ind_list))
    indices = _make_indices(ind_list.split(','), len(other_entities))
    values = _take(other_entities, indices)
    return values

#################################
## Model building transformers ##
#################################

# Each of the following handlers converts parsed fields into model instances.
# Each handler takes as input the dictionary of parsed fields and returns
# a modified dictionary containing model instances.
# The 'finished' field is a list of entity names that have already been
# converted. It is used by handlers to do work on entities they are related to.
# 
# !!! IMPORTANT !!!
# All handlers are *required* to maintain the order of entities in value fields.

def p_dict(d, level=0):
    keys = sorted(d.keys())
    for key in keys:
        if isinstance(d[key], dict):
            print ('\t' * level) + key + ' =>'
            p_dict(d[key], level + 1)
        else:
            print ('\t' * level) + '{0} => {1}'.format(key, d[key])

# Handles building entities upon which others depend.
def _pre_handle(parsed, wb, *entity_names):
    print 'Calling _pre_handle to handle ' + ' and '.join(entity_names)
    print 'Parsed["finished"] = ' + str(parsed['finished'])
    for entity in entity_names:
        if entity not in parsed['finished']:
            print 'Unfound ' + entity
            parsed = _post_collection_handlers[entity](parsed, wb)
            parsed['finished'].append(entity)
    print '_-_-_ End of pre handle _-_-_'
    #p_dict(parsed)
    return parsed

def _action_h(parsed, wb):
    print '### action_h'
    parsed = _pre_handle(parsed, wb, STATE_AUTHORITIES, INTERNATIONAL_AUTHORITIES, EVENTS)
    for i in range(len(parsed[ACTIONS])):
        action = entities.translate_fields(parsed[ACTIONS][i], entities.action)
        action['state_bodies_approached'] = _from_indices(
            action['state_bodies_approached'], parsed[STATE_AUTHORITIES])
        action['international_bodies_approached'] = _from_indices(
            action['international_bodies_approached'], parsed[INTERNATIONAL_AUTHORITIES])
        action['events'] = _from_indices(action['events'], parsed[EVENTS])
        parsed[ACTIONS][i] = Action(**action)
    return parsed

def _actor_h(parsed, wb):
    print '### actor_h'
    parsed = _pre_handle(parsed, wb, ORGANISATIONS, PROFESSIONS, LOCATIONS)
    for i in range(len(parsed[ACTORS])):
        actor = entities.translate_fields(parsed[ACTORS][i], entities.actor)
        actor['organisations'] = _from_indices(actor['organisations'], parsed[ORGANISATIONS])
        actor['professions'] = _from_indices(actor['professions'], parsed[PROFESSIONS])
        actor['locations'] = _from_indices(actor['locations'], parsed[LOCATIONS])
        actor['birth_date'] = _excel_parse_date(actor['birth_date'], wb)
        actor['gender'] = actor['gender'].lower() == MALE # Not trying to be sexist, just efficient.
        actor['is_activist'] = actor['is_activist'].lower() == YES
        parsed[ACTORS][i] = Actor(**actor)
    return parsed

def _event_h(parsed, wb):
    print '### event_h'
    parsed = _pre_handle(parsed, wb,
        RELEASE_TYPES, LOCATIONS, PRISONS, SOURCES, ACTORS, RIGHTS_VIOLATIONS)
    for i in range(len(parsed[EVENTS])):
        event = entities.translate_fields(parsed[EVENTS][i], entities.event)
        event['release_types'] = _from_indices(event['release_types'], parsed[RELEASE_TYPES])
        event['locations'] = _from_indices(event['locations'], parsed[LOCATIONS])
        event['prisons'] = _from_indices(event['prisons'], parsed[PRISONS])
        event['sources'] = _from_indices(event['sources'], parsed[SOURCES])
        event['witnesses'] = _from_indices(event['witnesses'], parsed[ACTORS])
        event['victims'] = _from_indices(event['victims'], parsed[ACTORS])
        event['perpetrators'] = _from_indices(event['perpetrators'], parsed[ACTORS])
        event['event_types'] = _from_indices(
            event['event_types'], parsed[RIGHTS_VIOLATIONS])
        event['event_start'] = _excel_parse_date(event['event_start'], wb)
        event['event_end'] = _excel_parse_date(event['event_end'], wb)
        event['report_date'] = _excel_parse_date(event['report_date'], wb)
        event['psych_assist'] = event['psych_assist'].lower() == YES
        event['material_assist'] = event['material_assist'].lower() == YES
        event['was_activist'] = event['was_activist'].lower() == YES
        event['victim_is_complainant'] = event['victim_is_complainant'].lower() == YES
        event['allow_storage'] = event['allow_storage'].lower() == YES
        event['allow_publishing'] = event['allow_publishing'].lower() == YES
        event['allow_representation'] = event['allow_representation'].lower() == YES
        event['data_is_sensitive'] = event['data_is_sensitive'].lower() == YES
        parsed[EVENTS][i] = Event(**event)
    return parsed

def _report_h(parsed, wb):
    print '### report_h'
    parsed = _pre_handle(parsed, wb, EVENTS)
    for i in range(len(parsed[REPORTS])):
        report = entities.translate_fields(parsed[REPORTS][i], entities.report)
        report['events'] = _from_indices(report['events'], parsed[EVENTS])
        parsed[REPORTS][i] = Report(**report)
    return parsed 

def _inter_auth_h(parsed, wb):
    print '### inter_auth_h'
    for i in range(len(parsed[INTERNATIONAL_AUTHORITIES])):
        ia = entities.translate_fields(
            parsed[INTERNATIONAL_AUTHORITIES][i], entities.international_authority)
        parsed[INTERNATIONAL_AUTHORITIES][i] = InternationalAuthority(**ia)
    return parsed

def _state_auth_h(parsed, wb):
    print '### state_auth_h'
    for i in range(len(parsed[STATE_AUTHORITIES])):
        sa = entities.translate_fields(parsed[STATE_AUTHORITIES][i], entities.state_authority)
        parsed[STATE_AUTHORITIES][i] = StateAuthority(**sa)
    return parsed

def _organisation_h(parsed, wb):
    print '### organisation_h'
    parsed = _pre_handle(parsed, wb, LOCATIONS)
    for i in range(len(parsed[ORGANISATIONS])):
        organisation = entities.translate_fields(parsed[ORGANISATIONS][i], entities.organisation)
        organisation['locations'] = _from_indices(organisation['locations'], parsed[LOCATIONS])
        parsed[ORGANISATIONS][i] = Organisation(**organisation)
    return parsed

def _profession_h(parsed, wb):
    print '### profession_h'
    for i in range(len(parsed[PROFESSIONS])):
        profession = entities.translate_fields(parsed[PROFESSIONS][i], entities.profession)
        parsed[PROFESSIONS][i] = Profession(**profession)
    return parsed

def _location_h(parsed, wb):
    print '### location_h'
    for i in range(len(parsed[LOCATIONS])):
        location = entities.translate_fields(parsed[LOCATIONS][i], entities.location)
        parsed[LOCATIONS][i] = Location(**location)
    return parsed

def _event_type_h(parsed, wb):
    print '### event_type_h'
    for i in range(len(parsed[RIGHTS_VIOLATIONS])):
        rv = entities.translate_fields(parsed[RIGHTS_VIOLATIONS][i], entities.event_type)
        parsed[RIGHTS_VIOLATIONS][i] = EventType(**rv)
    return parsed

def _release_type_h(parsed, wb):
    print '### release_type_h'
    for i in range(len(parsed[RELEASE_TYPES])):
        rt = entities.translate_fields(parsed[RELEASE_TYPES][i], entities.release_type)
        parsed[RELEASE_TYPES][i] = ReleaseType(**rt)
    return parsed

def _source_h(parsed, wb):
    print '### source_h'
    parsed = _pre_handle(parsed, wb, ORGANISATIONS)
    for i in range(len(parsed[SOURCES])):
        source = entities.translate_fields(parsed[SOURCES][i], entities.source)
        source['organisations'] = _from_indices(source['organisations'], parsed[ORGANISATIONS])
        parsed[SOURCES][i] = Source(**source)
    return parsed

def _prison_type_h(parsed, wb):
    print '### prison_type_h'
    for i in range(len(parsed[PRISON_TYPES])):
        pt = entities.translate_fields(parsed[PRISON_TYPES][i], entities.prison_type)
        parsed[PRISON_TYPES][i] = PrisonType(**pt)
    return parsed

def _prison_h(parsed, wb):
    print '### prison_h'
    parsed = _pre_handle(parsed, wb, LOCATIONS, PRISON_TYPES)
    for i in range(len(parsed[PRISONS])):
        prison = entities.translate_fields(parsed[PRISONS][i], entities.prison)
        prison['locations'] = _from_indices(prison['locations'], parsed[LOCATIONS])
        prison['prison_types'] = _from_indices(prison['prison_types'], parsed[PRISON_TYPES])
        parsed[PRISONS][i] = Prison(**prison)
    return parsed

# Mapping of entitiy names, as listed in the document template,
# to handler functions that will convert parsed data into usable
# model instances.
_post_collection_handlers = {
    ACTIONS: _action_h,
    ACTORS: _actor_h,
    EVENTS: _event_h,
    REPORTS: _report_h,
    INTERNATIONAL_AUTHORITIES: _inter_auth_h,
    STATE_AUTHORITIES: _state_auth_h,
    ORGANISATIONS: _organisation_h,
    PROFESSIONS: _profession_h,
    LOCATIONS: _location_h,
    RELEASE_TYPES: _release_type_h,
    RIGHTS_VIOLATIONS: _event_type_h,
    SOURCES: _source_h,
    PRISON_TYPES: _prison_type_h,
    PRISONS: _prison_h
}

# Options for existing parsers.
_parsers = {
    ORG1_DOCX: _parse_org1_docx,
    ORG2_DOCX: _parse_org2_docx,
    EXCEL_DOC: _parse_excel_template
}

def _parse_error(stream):
    return {'error': 'No parser exists for the provided filetype.'}

# Public

def parse(stream, filetype):
    '''Parse the contents of a report file with one of the accepted filetypes into a dict'''
    parse = _parsers.get(filetype, _parse_error)
    return parse(stream)
